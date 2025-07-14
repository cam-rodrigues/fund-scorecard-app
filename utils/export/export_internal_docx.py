from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

def export_internal_docx(df, proposal_html, file_obj):
    doc = Document()

    # === Title Section ===
    title = doc.add_heading("Internal Proposal", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("Detailed Analysis", style='Intense Quote')

    doc.add_paragraph()  # Spacer

    # === Parse Proposal Summary ===
    soup = BeautifulSoup(proposal_html, "html.parser")

    # Add Recommendation Heading
    doc.add_heading("Recommendation", level=2)

    # Extract and format parts of the proposal
    if soup.h3:
        doc.add_paragraph(soup.h3.get_text(), style='Heading 3')

    if soup.find('b'):
        doc.add_paragraph("Primary Candidate:", style='List Bullet')
        doc.add_paragraph(soup.find('b').get_text(), style='List Continue')

    if soup.find('em'):
        doc.add_paragraph(soup.find('em').get_text(), style='List Continue')

    # Add bullets from <ul><li>
    ul = soup.find('ul')
    if ul:
        for li in ul.find_all('li'):
            doc.add_paragraph(li.get_text(), style='List Bullet')

    doc.add_paragraph()  # Spacer

    # === Table Section ===
    doc.add_heading("Performance Table", level=2)

    if not df.empty:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Format header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                para = row_cells[i].paragraphs[0]
                run = para.add_run(str(val))
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === Export ===
    doc.save(file_obj)
