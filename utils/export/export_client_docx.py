from docx import Document

def export_client_docx(df, proposal_html, file_obj):
    """
    Creates a DOCX file for client presentation and saves it to the file_obj (a BytesIO buffer).
    """
    doc = Document()
    doc.add_heading("Client Proposal", level=1)

    # Add the proposal summary
    doc.add_paragraph("Proposal Summary")
    doc.add_paragraph(proposal_html)

    # Create a table of the fund data
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(file_obj)  # Save to the BytesIO buffer instead of a file path
