import streamlit as st
import pdfplumber
import pandas as pd
import re
import numpy as np
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def extract_fund_performance(pdf_file):
    data = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text or "Fund Performance: Current vs." not in text:
                continue
            lines = text.split("\n")
            fund_name = None
            for line in lines:
                if re.search(r"[A-Z]{4,6}X", line):
                    fund_name = line.strip()
                    continue
                if fund_name and re.search(r"[-+]?\d+\.\d+", line):
                    numbers = re.findall(r"[-+]?\d+\.\d+", line)
                    if len(numbers) >= 6:
                        data.append({
                            "Fund": fund_name,
                            "QTD": float(numbers[0]),
                            "YTD": float(numbers[1]),
                            "1 Yr": float(numbers[2]),
                            "3 Yr": float(numbers[3]),
                            "5 Yr": float(numbers[4]),
                            "10 Yr": float(numbers[5]),
                        })
                        fund_name = None
    return pd.DataFrame(data)

def generate_summary(df):
    if df.empty:
        return "No valid funds found."

    benchmark = {
        "Fund": "S&P 500 (Benchmark)",
        "QTD": 2.00, "YTD": 6.00, "1 Yr": 12.00, "3 Yr": 8.00, "5 Yr": 10.00, "10 Yr": 10.50
    }
    df = pd.concat([df, pd.DataFrame([benchmark])], ignore_index=True)

    trailing = ["QTD", "YTD", "1 Yr", "3 Yr", "5 Yr", "10 Yr"]
    others = df[df["Fund"] != "S&P 500 (Benchmark)"]
    benchmark_row = df[df["Fund"] == "S&P 500 (Benchmark)"].iloc[0]

    beat_counts = (others[trailing] > benchmark_row[trailing]).sum(axis=1)
    top_idx = beat_counts.idxmax()
    best = others.loc[top_idx]

    return f"""
Fund Summary:
- Total Funds Compared: {len(others)}
- Fund that outperformed S&P 500 the most: {best['Fund']} ({beat_counts[top_idx]} of 6 periods)
"""

def export_to_docx(summary_text):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Fund Performance Summary", level=1)
    doc.add_paragraph(summary_text)
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(summary_text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica", 12)
    for line in summary_text.strip().split("\n"):
        c.drawString(50, y, line.strip())
        y -= 20
        if y < 50:
            c.showPage()
            y = height - 50
    c.save()
    buffer.seek(0)
    return buffer

# ✅ MAIN FUNCTION for use in multipage apps
def run():
    st.set_page_config(page_title="Simple Fund Summary Export", layout="wide")
    st.title("📄 Fund Summary Extractor + Exporter")

    uploaded_pdf = st.file_uploader("Step 1: Upload MPI-style PDF", type=["pdf"])
    if uploaded_pdf:
        with st.spinner("Reading and extracting fund performance..."):
            df = extract_fund_performance(uploaded_pdf)

        if df.empty:
            st.error("No fund data found.")
        else:
            st.success(f"{len(df)} funds found.")
            st.dataframe(df)

            summary = generate_summary(df)
            st.markdown("### Step 2: Generated Summary")
            st.text_area("Summary Preview", value=summary, height=150)

            st.markdown("### Step 3: Export")
            file_type = st.radio("Choose format:", ["Word (.docx)", "PDF (.pdf)"])
            if st.button("📥 Export Summary"):
                if file_type == "Word (.docx)":
                    doc = export_to_docx(summary)
                    st.download_button("Download Word Doc", doc, "Fund_Summary.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    pdf = export_to_pdf(summary)
                    st.download_button("Download PDF", pdf, "Fund_Summary.pdf", mime="application/pdf")
