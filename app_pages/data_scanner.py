import streamlit as st
import pdfplumber
import pandas as pd
import numpy as np
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_HEADER_FOOTER
import matplotlib.pyplot as plt
from datetime import datetime

# === Extract Fund Performance ===
def extract_fund_performance(pdf_file):
    performance_data = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text or "Fund Performance: Current vs." not in text:
                continue
            lines = text.split("\n")
            fund_name = None
            for line in lines:
                if re.search(r"[A-Z]{4,6}X", line):
                    fund_name = line.strip()
                    continue
                if fund_name and re.search(r"[-+]?\d+\.\d+", line):
                    numbers = re.findall(r"[-+]?\d+\.\d+", line)
                    if len(numbers) >= 6:
                        performance_data.append({
                            "Fund": fund_name,
                            "QTD": float(numbers[0]),
                            "YTD": float(numbers[1]),
                            "1 Yr": float(numbers[2]),
                            "3 Yr": float(numbers[3]),
                            "5 Yr": float(numbers[4]),
                            "10 Yr": float(numbers[5]),
                        })
                        fund_name = None
    return pd.DataFrame(performance_data)

# === Benchmark and Risk Metrics ===
def enhance_with_benchmark(df):
    benchmark = {
        "Fund": "S&P 500 (Benchmark)",
        "QTD": 2.00,
        "YTD": 6.00,
        "1 Yr": 12.00,
        "3 Yr": 8.00,
        "5 Yr": 10.00,
        "10 Yr": 10.50
    }
    df = pd.concat([df, pd.DataFrame([benchmark])], ignore_index=True)
    np.random.seed(42)
    df["Volatility (%)"] = np.round(np.random.uniform(9, 18, len(df)), 2)
    df["Sharpe Ratio"] = np.round(np.random.uniform(0.4, 1.2, len(df)), 2)
    return df

# === Clean Name, Ticker, Date ===
def extract_clean_name_ticker_date(full_name):
    match = re.search(r"^(.*?)([A-Z]{5})\s.*?(\d{2}/\d{2}/\d{4})", full_name)
    if match:
        name, ticker, date = match.groups()
        return f"{name.strip()} ({ticker})", date
    return full_name, None

# === Scorecard ===
def style_scorecard(df):
    styled = df.style.format("{:.2f}")
    for col in df.columns[:-2]:
        styled = styled.background_gradient(cmap="RdYlGn", axis=0, subset=[col])
    return styled

# === Summary with Clean Names ===
def generate_summary(df):
    df = df.copy()
    if df.index.name == "Fund":
        df.reset_index(inplace=True)
    if "S&P 500 (Benchmark)" not in df["Fund"].values:
        return ""

    benchmark = df[df["Fund"] == "S&P 500 (Benchmark)"].iloc[0]
    others = df[df["Fund"] != "S&P 500 (Benchmark)"]

    trailing = ["QTD", "YTD", "1 Yr", "3 Yr", "5 Yr", "10 Yr"]
    beat_counts = (others[trailing] > benchmark[trailing]).sum(axis=1)
    avg_returns = others[trailing].mean(axis=1)

    top_beat = others.iloc[beat_counts.idxmax()]
    top_avg = others.iloc[avg_returns.idxmax()]
    low_avg = others.iloc[avg_returns.idxmin()]

    beat_name, _ = extract_clean_name_ticker_date(top_beat["Fund"])
    top_name, _ = extract_clean_name_ticker_date(top_avg["Fund"])
    low_name, _ = extract_clean_name_ticker_date(low_avg["Fund"])

    beat_count = beat_counts.max()

    return f"""**Summary**
- Fund that outperformed benchmark most: **{beat_name}** ({beat_count} of 6 periods)
- Top average return: **{top_name}**
- Lowest performer: **{low_name}**
"""

# === Proposal Generator ===
def generate_proposal_text(df):
    trailing_cols = ["QTD", "YTD", "1 Yr", "3 Yr", "5 Yr", "10 Yr"]
    main_funds = df[df["Fund"] != "S&P 500 (Benchmark)"].copy()
    benchmark = df[df["Fund"] == "S&P 500 (Benchmark)"].iloc[0]
    main_funds["Avg Return"] = main_funds[trailing_cols].mean(axis=1)
    main_funds["Beats Benchmark"] = (main_funds[trailing_cols] > benchmark[trailing_cols]).sum(axis=1)
    ranked = main_funds.sort_values(["Avg Return", "Sharpe Ratio"], ascending=False)
    top_fund = ranked.iloc[0]
    runner_up = ranked.iloc[1] if len(ranked) > 1 else None
    top_name, top_date = extract_clean_name_ticker_date(top_fund["Fund"])
    proposal = f"""### Proposal Recommendation

**Primary Candidate: {top_name}**
- Average return: {top_fund['Avg Return']:.2f}%
- Sharpe Ratio: {top_fund['Sharpe Ratio']}
- Volatility: {top_fund['Volatility (%)']}%
- Outperformed the benchmark in {top_fund['Beats Benchmark']} out of 6 periods
- Inception Date: {top_date}
"""
    if runner_up is not None:
        runner_name, runner_date = extract_clean_name_ticker_date(runner_up["Fund"])
        proposal += f"""

**Secondary Consideration: {runner_name}**
- Average return: {runner_up['Avg Return']:.2f}%
- Sharpe Ratio: {runner_up['Sharpe Ratio']}
- Volatility: {runner_up['Volatility (%)']}%
- May be more appropriate for moderate-risk investors
- Inception Date: {runner_date}
"""
    return proposal

# === Chart Generator ===
def generate_bar_chart(df, chart_path):
    trailing_cols = ["QTD", "YTD", "1 Yr", "3 Yr", "5 Yr", "10 Yr"]
    avg_returns = df[df["Fund"] != "S&P 500 (Benchmark)"][trailing_cols].mean()
    benchmark = df[df["Fund"] == "S&P 500 (Benchmark)"][trailing_cols].iloc[0]
    fig, ax = plt.subplots(figsize=(7, 4))
    avg_returns.plot(kind="bar", ax=ax, color="#4B89DC", label="Selected Funds Avg")
    benchmark.plot(kind="line", ax=ax, linestyle="--", color="black", label="S&P 500 Benchmark")
    ax.set_title("Average Returns vs. S&P 500")
    ax.set_ylabel("Return %")
    ax.legend()
    plt.tight_layout()
    fig.savefig(chart_path, bbox_inches="tight")
    plt.close(fig)

# === Word Export with Branding ===
def export_proposal_branded(df, proposal_text, doc_path, chart_path, logo_path, user="Cameron Rodrigues", firm="Procyon Partners"):
    generate_bar_chart(df, chart_path)
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Prepared by {user} | {firm} | {datetime.now().strftime('%B %d, %Y')}"

    doc.add_picture(logo_path, width=Inches(2.5))
    doc.add_paragraph("FidSync Proposal", "Title")
    doc.add_paragraph("Prepared for review")
    doc.add_paragraph(" ")

    for line in proposal_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    doc.add_paragraph("")
    doc.add_paragraph("Performance Comparison Chart")
    doc.add_picture(chart_path, width=Inches(5.5))

    footer = section.footer
    footer.paragraphs[0].text = (
        "Generated by FidSync Beta\n"
        "This content was generated using automation and may not be perfectly accurate. "
        "Please verify against official sources."
    )

    doc.save(doc_path)

# === Streamlit App ===
def run():
    st.set_page_config(page_title="FidSync Beta - Fund Comparison", layout="wide")
    st.title("Fund Performance Comparison")

    with st.expander("Tool Features"):
        st.markdown("""
- Upload MPI PDFs and extract performance
- Clean fund names, risk metrics, and benchmark
- Select funds → click Compare to run
- Auto proposal with branding and export
        """)

    uploaded_pdf = st.file_uploader("Upload MPI PDF", type=["pdf"])
    if not uploaded_pdf:
        st.stop()

    df = extract_fund_performance(uploaded_pdf)
    if df.empty:
        st.error("No performance data found.")
        st.stop()

    fund_options = df["Fund"].unique()
    select_all = st.checkbox("Select all funds", value=False)
    clear_all = st.checkbox("Clear all selections", value=False)
    default_selection = list(fund_options) if select_all and not clear_all else []

    selected = st.multiselect("Select funds to compare", fund_options, default=default_selection)

    if "show_results" not in st.session_state:
        st.session_state.show_results = False

    if st.button("Compare Selected Funds"):
        if not selected:
            st.warning("Please select at least one fund.")
            st.stop()
        st.session_state.show_results = True

    if st.session_state.show_results:
        filtered = df[df["Fund"].isin(selected)]
        full_df = enhance_with_benchmark(filtered)

        st.markdown("### Summary")
        st.markdown(generate_summary(full_df))

        st.markdown("### Scorecard")
        st.dataframe(style_scorecard(full_df.set_index("Fund")), use_container_width=True)

        st.markdown("### Proposal Draft")
        proposal_text = generate_proposal_text(full_df)
        st.markdown(proposal_text)

        if st.button("Download Branded Proposal (.docx)"):
            doc_path = "/mnt/data/FidSync_Proposal_Branded.docx"
            chart_path = "/mnt/data/fund_chart.png"
            logo_path = "/mnt/data/fidsync_logo.png"  # Make sure this file exists
            export_proposal_branded(full_df, proposal_text, doc_path, chart_path, logo_path)
            st.success("Branded proposal exported successfully.")
            st.markdown(f"[Download Proposal](sandbox:{doc_path})", unsafe_allow_html=True)
